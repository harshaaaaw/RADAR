import random
import string
import zipfile
import shutil
from pathlib import Path

# Libraries we verified existence of
from PIL import Image, ImageDraw
from docx import Document
from docx.shared import Inches

TARGET_DIR = Path(r"C:\Users\DELL\Downloads\TestDocuments")
TEMP_IMG_DIR = TARGET_DIR / "temp_assets"

def setup():
    TARGET_DIR.mkdir(parents=True, exist_ok=True)
    TEMP_IMG_DIR.mkdir(parents=True, exist_ok=True)
    print(f"Generating files in {TARGET_DIR}...")

def generate_random_text(lines=10):
    text = []
    for _ in range(lines):
        sentence = ' '.join(''.join(random.choices(string.ascii_lowercase, k=random.randint(3, 10))) for _ in range(random.randint(5, 15)))
        text.append(sentence.capitalize() + '.')
    return '\n'.join(text)

def generate_base_images():
    print("Generating base images...")
    images = []
    for i in range(5):
        img_path = TEMP_IMG_DIR / f"base_img_{i}.png"
        img = Image.new('RGB', (400, 200), color=(random.randint(0, 255), random.randint(0, 255), random.randint(0, 255)))
        d = ImageDraw.Draw(img)
        d.text((10, 80), f"TEST OCR CONTENT {i}\nRandom: {random.randint(1000,9999)}", fill=(255, 255, 255))
        img.save(img_path)
        images.append(img_path)
    return images

def generate_text_files(count=1000):
    print(f"Generating {count} text files...")
    for i in range(count):
        with open(TARGET_DIR / f"stress_txt_{i}.txt", 'w', encoding='utf-8') as f:
            f.write(f"File ID: {i}\n")
            f.write(generate_random_text(random.randint(5, 50)))
        if i % 100 == 0: print(f"  {i}/{count}")

def generate_docx_files(count=400, base_images=None):
    if base_images is None:
        base_images = []
    print(f"Generating {count} Word docs...")
    for i in range(count):
        doc = Document()
        doc.add_heading(f'Stress Test Document {i}', 0)
        doc.add_paragraph(f'This is a generated document used for system stress testing. ID: {i}')
        doc.add_paragraph(generate_random_text(5))
        
        # Embed image occasionally (50% chance)
        if base_images and random.random() > 0.5:
            img = random.choice(base_images)
            doc.add_picture(str(img), width=Inches(4.0))
            doc.add_paragraph("Caption: Embedded Image for OCR test.")
            
        doc.save(TARGET_DIR / f"stress_doc_{i}.docx")
        if i % 50 == 0: print(f"  {i}/{count}")

def generate_image_files(count=200):
    print(f"Generating {count} image files...")
    for i in range(count):
        img = Image.new('RGB', (800, 600), color='white')
        d = ImageDraw.Draw(img)
        d.text((50, 50), f"DIRECT IMAGE OCR TEST {i}\nConfidential Data: {random.randint(10000,99999)}", fill='black')
        img.save(TARGET_DIR / f"stress_img_{i}.png")
        if i % 50 == 0: print(f"  {i}/{count}")

def generate_zips(count=200):
    print(f"Generating {count} Zip archives...")
    for i in range(count):
        zip_path = TARGET_DIR / f"stress_zip_{i}.zip"
        with zipfile.ZipFile(zip_path, 'w') as zf:
            zf.writestr('inner_text.txt', f"Content inside zip {i}")
            # Add an image inside zip
            img_name = f"inner_img_{i}.png"
            img_path = TEMP_IMG_DIR / "base_img_0.png"
            if img_path.exists():
                zf.write(img_path, img_name)
        if i % 50 == 0: print(f"  {i}/{count}")

def generate_large_files(count=10):
    print(f"Generating {count} Large files...")
    for i in range(count):
        with open(TARGET_DIR / f"stress_large_{i}.txt", 'w', encoding='utf-8') as f:
            # Write 5MB of data
            f.write("LARGE FILE START\n")
            f.write("A" * (5 * 1024 * 1024))
            f.write("\nLARGE FILE END")

def generate_corrupt_files(count=50):
    print(f"Generating {count} Corrupt/Deep files...")
    for i in range(count):
        # Empty
        with open(TARGET_DIR / f"stress_corrupt_{i}.docx", 'wb') as f:
            f.write(b'PK_FAKE_HEADER_BUT_BROKEN')

def main():
    setup()
    base_images = generate_base_images()
    
    generate_text_files(1000)
    generate_docx_files(400, base_images)
    generate_image_files(200)
    generate_zips(200)
    generate_large_files(10)
    generate_corrupt_files(50)
    generate_challenging_files(50)
    
    # Cleanup temp
    try:
        shutil.rmtree(TEMP_IMG_DIR)
    except:
        pass
        
    print("\nGeneration Complete!")
    print(f"Total files in {TARGET_DIR}: {len(list(TARGET_DIR.glob('*')))}")

def generate_challenging_files(count=50):
    print(f"Generating {count} Challenging files (Rotated, Inverted, PDF)...")
    import numpy as np
    import cv2
    
    font = cv2.FONT_HERSHEY_SIMPLEX
    
    for i in range(count):
        # 1. Rotated Images
        img = np.ones((600, 800), dtype=np.uint8) * 255
        cv2.putText(img, f"ROTATION TEST {i}", (100, 300), font, 2, (0), 3)
        cv2.putText(img, f"Secret Key: ROT_{random.randint(1000,9999)}", (100, 400), font, 1, (0), 2)
        
        # Rotate
        angle = random.choice([90, 180, 270])
        if angle == 90:
            rotated = cv2.rotate(img, cv2.ROTATE_90_CLOCKWISE)
        elif angle == 180:
            rotated = cv2.rotate(img, cv2.ROTATE_180)
        elif angle == 270:
            rotated = cv2.rotate(img, cv2.ROTATE_90_COUNTERCLOCKWISE)
        
        cv2.imwrite(str(TARGET_DIR / f"stress_challenging_rot{angle}_{i}.png"), rotated)
        
        # 2. Inverted Images (White on Black)
        img_inv = np.ones((200, 600), dtype=np.uint8) * 0 # Black
        cv2.putText(img_inv, f"INVERTED TEST {i}", (50, 100), font, 1.5, (255), 2) # White text
        cv2.putText(img_inv, f"Inverted Key: INV_{random.randint(1000,9999)}", (50, 150), font, 1, (255), 2)
        cv2.imwrite(str(TARGET_DIR / f"stress_challenging_inv_{i}.png"), img_inv)
        
        # 3. Shadowed/Noisy Images
        img_shd = np.ones((400, 600), dtype=np.uint8) * 255
        cv2.putText(img_shd, f"SHADOW TEST {i}", (50, 100), font, 1.5, (0), 2)
        cv2.putText(img_shd, f"Shadow Key: SHD_{random.randint(1000,9999)}", (50, 200), font, 1, (0), 2)
        
        # Add gradient shadow
        rows, cols = img_shd.shape
        shadow = np.zeros((rows, cols), dtype=np.uint8)
        for c in range(cols):
            shadow[:, c] = int(255 * (c / cols) * 0.7) # Darker on right
        
        img_shd = img_shd.astype(np.int16) - shadow
        img_shd = np.clip(img_shd, 0, 255).astype(np.uint8)
        cv2.imwrite(str(TARGET_DIR / f"stress_challenging_shadow_{i}.png"), img_shd)
        
        # 4. Standard PDF with Image (using PIL)
        # Create an image first
        pil_img = Image.new('RGB', (800, 1000), color='white')
        d = ImageDraw.Draw(pil_img)
        d.text((50, 50), f"PDF IMAGE TEST {i}\nPDF Key: PDF_{random.randint(1000,9999)}", fill='black')
        pil_img.save(TARGET_DIR / f"stress_challenging_pdf_{i}.pdf", "PDF", resolution=100.0)


if __name__ == "__main__":
    main()
