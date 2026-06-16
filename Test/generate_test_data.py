#!/usr/bin/env python3
"""
Test Data Generator for Enterprise Document Search System
Generates 500+ test documents in various formats to test all system capabilities
"""

import random
from pathlib import Path
from datetime import datetime, timedelta
import json

# Try to import optional libraries
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Install with: pip install python-docx")

try:
    from openpyxl import Workbook
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False
    print("Warning: openpyxl not installed. Install with: pip install openpyxl")

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("Warning: reportlab not installed. Install with: pip install reportlab")

try:
    import csv
    CSV_AVAILABLE = True
except ImportError:
    CSV_AVAILABLE = False


# Sample content templates
BUSINESS_TOPICS = [
    "Quarterly Financial Report",
    "Marketing Strategy",
    "Product Development Roadmap",
    "Customer Satisfaction Survey",
    "Employee Performance Review",
    "Sales Pipeline Analysis",
    "Project Status Update",
    "Risk Assessment Report",
    "Compliance Documentation",
    "Training Materials",
    "Meeting Minutes",
    "Budget Proposal",
    "Contract Agreement",
    "Technical Specification",
    "User Manual",
    "Policy Document",
    "Incident Report",
    "Quality Assurance Report",
    "Research Findings",
    "Strategic Planning"
]

DEPARTMENTS = [
    "Finance", "Marketing", "Sales", "HR", "IT", "Operations",
    "Legal", "R&D", "Customer Service", "Engineering"
]

SAMPLE_PARAGRAPHS = [
    "This document outlines the key findings from our recent analysis. The data shows significant growth in the target market segment, with a year-over-year increase of approximately 15%. Our team has identified several opportunities for expansion and improvement.",
    
    "The strategic objectives for this quarter include enhancing customer engagement, streamlining operational processes, and implementing new technology solutions. Each initiative has been assigned to specific team members with clear deliverables and timelines.",
    
    "Based on comprehensive market research, we recommend focusing on digital transformation initiatives. This approach will enable us to better serve our customers while reducing operational costs. The projected ROI is estimated at 25% within the first year.",
    
    "Our analysis indicates that customer satisfaction has improved by 12% compared to the previous period. Key factors contributing to this improvement include faster response times, enhanced product quality, and improved communication channels.",
    
    "The project team has successfully completed the initial phase ahead of schedule. All major milestones have been achieved, and stakeholder feedback has been overwhelmingly positive. We are now proceeding to the implementation phase.",
    
    "Risk mitigation strategies have been developed to address potential challenges. These include contingency planning, resource allocation adjustments, and enhanced monitoring protocols. Regular reviews will ensure timely identification of any issues.",
    
    "Training programs have been designed to upskill our workforce in emerging technologies. The curriculum covers both technical and soft skills, with a focus on practical application. Completion rates have exceeded expectations.",
    
    "Quality metrics demonstrate consistent improvement across all key performance indicators. Process optimization initiatives have resulted in reduced error rates and increased efficiency. Continuous improvement remains a top priority.",
    
    "Compliance requirements have been thoroughly reviewed and updated to reflect current regulations. All necessary documentation has been prepared and submitted to relevant authorities. Regular audits will ensure ongoing adherence.",
    
    "Innovation initiatives are driving competitive advantage in the marketplace. Our R&D team has developed several promising prototypes that address customer pain points. Patent applications are currently in progress."
]


def generate_random_text(num_paragraphs=3):
    """Generate random business-like text"""
    paragraphs = random.sample(SAMPLE_PARAGRAPHS, min(num_paragraphs, len(SAMPLE_PARAGRAPHS)))
    return "\n\n".join(paragraphs)


def generate_filename(topic, dept, index, extension):
    """Generate a realistic filename"""
    date_str = (datetime.now() - timedelta(days=random.randint(0, 365))).strftime("%Y%m%d")
    clean_topic = topic.replace(" ", "_")
    return f"{dept}_{clean_topic}_{date_str}_{index:04d}.{extension}"


def create_text_file(output_dir, filename, content):
    """Create a plain text file"""
    filepath = output_dir / filename
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(content)
    return filepath


def create_markdown_file(output_dir, filename, topic, content):
    """Create a Markdown file"""
    filepath = output_dir / filename
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(f"# {topic}\n\n")
        f.write(f"**Date:** {datetime.now().strftime('%Y-%m-%d')}\n\n")
        f.write(f"## Overview\n\n{content}\n\n")
        f.write("## Conclusion\n\nThis document provides important information for stakeholders.\n")
    return filepath


def create_html_file(output_dir, filename, topic, content):
    """Create an HTML file"""
    filepath = output_dir / filename
    html_content = f"""<!DOCTYPE html>
<html>
<head>
    <title>{topic}</title>
    <meta charset="UTF-8">
</head>
<body>
    <h1>{topic}</h1>
    <p><strong>Date:</strong> {datetime.now().strftime('%Y-%m-%d')}</p>
    <div>
        {content.replace(chr(10), '<br>')}
    </div>
    <footer>
        <p>Generated for testing purposes</p>
    </footer>
</body>
</html>"""
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(html_content)
    return filepath


def create_json_file(output_dir, filename, topic, dept):
    """Create a JSON file"""
    filepath = output_dir / filename
    data = {
        "document_type": topic,
        "department": dept,
        "created_date": datetime.now().isoformat(),
        "author": f"{random.choice(['John', 'Jane', 'Mike', 'Sarah'])} {random.choice(['Smith', 'Johnson', 'Williams', 'Brown'])}",
        "status": random.choice(["Draft", "Final", "Review", "Approved"]),
        "priority": random.choice(["High", "Medium", "Low"]),
        "content": generate_random_text(2),
        "metadata": {
            "version": f"{random.randint(1, 5)}.{random.randint(0, 9)}",
            "tags": random.sample(["important", "urgent", "confidential", "public", "internal"], k=2)
        }
    }
    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2)
    return filepath


def create_xml_file(output_dir, filename, topic, dept, content):
    """Create an XML file"""
    filepath = output_dir / filename
    xml_content = f"""<?xml version="1.0" encoding="UTF-8"?>
<document>
    <metadata>
        <title>{topic}</title>
        <department>{dept}</department>
        <date>{datetime.now().strftime('%Y-%m-%d')}</date>
        <author>Test Author</author>
    </metadata>
    <content>
        <section>
            <heading>Main Content</heading>
            <text>{content}</text>
        </section>
    </content>
</document>"""
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(xml_content)
    return filepath


def create_csv_file(output_dir, filename):
    """Create a CSV file"""
    if not CSV_AVAILABLE:
        return None
    
    filepath = output_dir / filename
    with open(filepath, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow(['ID', 'Name', 'Department', 'Date', 'Amount', 'Status'])
        for i in range(random.randint(10, 50)):
            writer.writerow([
                i + 1,
                f"Item {i + 1}",
                random.choice(DEPARTMENTS),
                (datetime.now() - timedelta(days=random.randint(0, 30))).strftime('%Y-%m-%d'),
                f"${random.randint(100, 10000)}",
                random.choice(['Pending', 'Approved', 'Completed'])
            ])
    return filepath


def create_docx_file(output_dir, filename, topic, content):
    """Create a Word document"""
    if not DOCX_AVAILABLE:
        return None
    
    filepath = output_dir / filename
    doc = Document()
    doc.add_heading(topic, 0)
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_heading('Overview', level=1)
    doc.add_paragraph(content)
    doc.add_heading('Summary', level=1)
    doc.add_paragraph('This document contains important information for review.')
    doc.save(str(filepath))
    return filepath


def create_xlsx_file(output_dir, filename):
    """Create an Excel file"""
    if not XLSX_AVAILABLE:
        return None
    
    filepath = output_dir / filename
    wb = Workbook()
    ws = wb.active
    ws.title = "Data"
    
    # Headers
    ws.append(['ID', 'Product', 'Category', 'Price', 'Quantity', 'Total'])
    
    # Data
    for i in range(random.randint(20, 100)):
        price = random.randint(10, 1000)
        qty = random.randint(1, 100)
        ws.append([
            i + 1,
            f"Product {i + 1}",
            random.choice(['Electronics', 'Furniture', 'Supplies', 'Equipment']),
            price,
            qty,
            price * qty
        ])
    
    wb.save(str(filepath))
    return filepath


def create_pdf_file(output_dir, filename, topic, content):
    """Create a PDF file"""
    if not PDF_AVAILABLE:
        return None
    
    filepath = output_dir / filename
    c = canvas.Canvas(str(filepath), pagesize=letter)
    width, height = letter
    
    # Title
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, height - 50, topic)
    
    # Date
    c.setFont("Helvetica", 10)
    c.drawString(50, height - 70, f"Date: {datetime.now().strftime('%Y-%m-%d')}")
    
    # Content
    c.setFont("Helvetica", 12)
    y_position = height - 100
    
    # Split content into lines
    words = content.split()
    line = ""
    for word in words:
        if len(line + word) < 80:
            line += word + " "
        else:
            c.drawString(50, y_position, line)
            y_position -= 15
            line = word + " "
            if y_position < 50:
                c.showPage()
                y_position = height - 50
    
    if line:
        c.drawString(50, y_position, line)
    
    c.save()
    return filepath


def generate_test_data(output_dir, num_files=500):
    """Generate test data files"""
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    
    print(f"Generating {num_files} test files in {output_dir}...")
    print("=" * 70)
    
    file_generators = []
    
    # Always available formats
    file_generators.extend([
        ('txt', create_text_file, 0.15),
        ('md', create_markdown_file, 0.10),
        ('html', create_html_file, 0.10),
        ('json', create_json_file, 0.10),
        ('xml', create_xml_file, 0.10),
    ])
    
    # Optional formats
    if CSV_AVAILABLE:
        file_generators.append(('csv', create_csv_file, 0.10))
    if DOCX_AVAILABLE:
        file_generators.append(('docx', create_docx_file, 0.15))
    if XLSX_AVAILABLE:
        file_generators.append(('xlsx', create_xlsx_file, 0.10))
    if PDF_AVAILABLE:
        file_generators.append(('pdf', create_pdf_file, 0.10))
    
    # Normalize weights
    total_weight = sum(w for _, _, w in file_generators)
    file_generators = [(ext, func, w/total_weight) for ext, func, w in file_generators]
    
    stats = {ext: 0 for ext, _, _ in file_generators}
    created_files = []
    
    for i in range(num_files):
        # Select format based on weights
        rand = random.random()
        cumulative = 0
        selected_ext = None
        selected_func = None
        
        for ext, func, weight in file_generators:
            cumulative += weight
            if rand <= cumulative:
                selected_ext = ext
                selected_func = func
                break
        
        if not selected_ext:
            selected_ext, selected_func, _ = file_generators[0]
        
        # Generate file
        topic = random.choice(BUSINESS_TOPICS)
        dept = random.choice(DEPARTMENTS)
        content = generate_random_text(random.randint(2, 5))
        filename = generate_filename(topic, dept, i, selected_ext)
        
        try:
            if selected_ext == 'txt':
                filepath = selected_func(output_path, filename, f"{topic}\n\n{content}")
            elif selected_ext == 'md':
                filepath = selected_func(output_path, filename, topic, content)
            elif selected_ext == 'html':
                filepath = selected_func(output_path, filename, topic, content)
            elif selected_ext == 'json':
                filepath = selected_func(output_path, filename, topic, dept)
            elif selected_ext == 'xml':
                filepath = selected_func(output_path, filename, topic, dept, content)
            elif selected_ext == 'csv':
                filepath = selected_func(output_path, filename)
            elif selected_ext == 'docx':
                filepath = selected_func(output_path, filename, topic, content)
            elif selected_ext == 'xlsx':
                filepath = selected_func(output_path, filename)
            elif selected_ext == 'pdf':
                filepath = selected_func(output_path, filename, topic, content)
            
            if filepath:
                stats[selected_ext] += 1
                created_files.append(filepath)
                
                if (i + 1) % 50 == 0:
                    print(f"Progress: {i + 1}/{num_files} files created...")
        
        except Exception as e:
            print(f"Error creating {filename}: {e}")
    
    print("=" * 70)
    print(f"\nGeneration complete! Created {len(created_files)} files")
    print("\nFile type distribution:")
    for ext, count in sorted(stats.items()):
        print(f"  {ext.upper():6s}: {count:4d} files")
    
    print(f"\nFiles created in: {output_path.absolute()}")
    print("\nMissing libraries (optional):")
    if not DOCX_AVAILABLE:
        print("  - python-docx (for .docx files): pip install python-docx")
    if not XLSX_AVAILABLE:
        print("  - openpyxl (for .xlsx files): pip install openpyxl")
    if not PDF_AVAILABLE:
        print("  - reportlab (for .pdf files): pip install reportlab")
    
    return created_files


if __name__ == "__main__":
    import sys
    
    # Default output directory
    default_dir = r"C:\Users\DELL\Downloads\DocumentSearch\test_data"
    
    output_dir = sys.argv[1] if len(sys.argv) > 1 else default_dir
    num_files = int(sys.argv[2]) if len(sys.argv) > 2 else 500
    
    print("\n" + "=" * 70)
    print("Enterprise Document Search - Test Data Generator")
    print("=" * 70)
    print(f"Output directory: {output_dir}")
    print(f"Number of files: {num_files}")
    print("=" * 70 + "\n")
    
    generate_test_data(output_dir, num_files)
    
    print("\n" + "=" * 70)
    print("Ready to test! The document search system will now process these files.")
    print("=" * 70 + "\n")
