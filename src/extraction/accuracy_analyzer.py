"""
Accuracy Analyzer — Measures extraction accuracy per document.

Tier 1: OpenCV + Tesseract TSV (always available)
Tier 2: + YOLOv8n object detection (optional)
Tier 3: + DocTR ground-truth OCR (optional)

Outputs per-document:
  - extraction_accuracy (0-100%)
  - text_area_pct / non_text_area_pct
  - accuracy_loss_json (breakdown of WHERE accuracy was lost)
  - pipeline_type (text_extraction | ocr)
  - raw_char_count / processed_char_count
  - preprocessing_gain_pct
  - page_metrics_json (per-page breakdown)
  - accuracy_tier (tier1 | tier2 | tier3)
"""

from __future__ import annotations

import io
import json
import os
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import cv2
import numpy as np
from PIL import Image

from core.logging_manager import get_logger

logger = get_logger("extraction.accuracy")

# ---------- Lazy optional imports ----------


try:
    from ultralytics import YOLO
    _HAS_YOLO = True
except ImportError:
    _HAS_YOLO = False

try:
    from doctr.io import DocumentFile
    from doctr.models import ocr_predictor
    _HAS_DOCTR = True
except ImportError:
    _HAS_DOCTR = False

try:
    import pytesseract
    _HAS_PYTESSERACT = True
except ImportError:
    _HAS_PYTESSERACT = False


def _empty_metrics(pipeline_type: str = "text_extraction", tier: str = "tier1") -> Dict[str, Any]:
    """Return a metrics dict with all fields defaulted."""
    return {
        "pipeline_type": pipeline_type,
        "extraction_accuracy": None,
        "text_area_pct": None,
        "non_text_area_pct": None,
        "raw_char_count": None,
        "processed_char_count": None,
        "preprocessing_gain_pct": None,
        "accuracy_loss_json": "{}",
        "page_metrics_json": "[]",
        "accuracy_tier": tier,
    }


# =============================================================================
# AccuracyAnalyzer
# =============================================================================
class AccuracyAnalyzer:
    """Tiered document extraction accuracy analyzer."""

    def __init__(self, enable_yolo: bool = True, enable_doctr: bool = True):
        self.tier = "tier1"

        # Tier 2: YOLOv9
        self._yolo = None
        if enable_yolo and _HAS_YOLO:
            try:
                # Use a pre-trained YOLO model for document layout
                # We'll use yolov8n (nano) for CPU efficiency
                self._yolo = YOLO("yolov8n.pt")
                self.tier = "tier2"
                logger.info("AccuracyAnalyzer: YOLOv8n loaded (Tier 2)")
            except Exception as exc:
                logger.warning("AccuracyAnalyzer: YOLO load failed — staying Tier 1: %s", exc)

        # Tier 3: DocTR
        self._doctr = None
        if enable_doctr and _HAS_DOCTR:
            try:
                self._doctr = ocr_predictor(
                    det_arch="db_mobilenet_v3_large",
                    reco_arch="crnn_mobilenet_v3_small",
                    pretrained=True,
                )
                if self.tier == "tier2":
                    self.tier = "tier3"
                # else: keep current tier — DocTR alone without YOLO doesn't qualify for tier3
                logger.info("AccuracyAnalyzer: DocTR loaded (Tier 3)")
            except Exception as exc:
                logger.warning("AccuracyAnalyzer: DocTR load failed: %s", exc)

        logger.info("AccuracyAnalyzer initialized — active tier: %s", self.tier)

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------
    def analyze(
        self,
        file_path: str,
        extracted_text: str,
        tika_response: Optional[Dict] = None,
    ) -> Dict[str, Any]:
        """Analyze extraction accuracy for a non-OCR file.

        Routes to the appropriate per-format analyzer.
        """
        ext = Path(file_path).suffix.lower()
        try:
            if ext in {".txt", ".csv", ".md"}:
                return self._analyze_pure_text(file_path, extracted_text, ext)
            elif ext in {".json", ".xml"}:
                return self._analyze_structured_text(file_path, extracted_text, ext)
            elif ext == ".html":
                return self._analyze_html(file_path, extracted_text)
            elif ext == ".docx":
                return self._analyze_docx(file_path, extracted_text)
            elif ext == ".xlsx":
                return self._analyze_xlsx(file_path, extracted_text)
            elif ext == ".pdf":
                return self._analyze_text_pdf(file_path, extracted_text, tika_response)
            else:
                return self._analyze_generic(file_path, extracted_text, ext)
        except Exception as exc:
            logger.warning("Accuracy analysis failed for %s: %s", file_path, exc)
            m = _empty_metrics("text_extraction", self.tier)
            m["accuracy_loss_json"] = json.dumps({"error": str(exc)})
            return m

    def analyze_ocr_page(
        self,
        raw_image_bytes: bytes,
        preprocessed_image_bytes: bytes,
    ) -> Dict[str, Any]:
        """Analyze OCR accuracy for a single page image.

        Args:
            raw_image_bytes: Original image before any preprocessing.
            preprocessed_image_bytes: Image after best preprocessing strategy.

        Returns:
            Per-page accuracy metrics dict.
        """
        try:
            # Step 1: Zone segmentation
            zone_metrics = self._segment_page_opencv(raw_image_bytes)
            zone_metrics = self._refine_zones_yolo(raw_image_bytes, zone_metrics)

            # Step 2: Tesseract TSV char-level comparison
            tess_metrics = self._measure_tesseract_coverage(
                raw_image_bytes, preprocessed_image_bytes
            )

            # Step 3: DocTR ground truth (Tier 3 only)
            doctr_metrics = None
            if self._doctr:
                doctr_metrics = self._measure_doctr_ground_truth(
                    raw_image_bytes, preprocessed_image_bytes
                )

            # Step 4: Estimate total chars in text zone
            estimated_total = self._estimate_total_chars(
                raw_image_bytes, zone_metrics["text_area_pct"]
            )

            # Step 5: Build loss breakdown
            loss = self._build_loss_breakdown(
                zone_metrics, tess_metrics, doctr_metrics, estimated_total
            )

            return {
                "pipeline_type": "ocr",
                "extraction_accuracy": loss["extraction_accuracy"],
                "text_area_pct": zone_metrics["text_area_pct"],
                "non_text_area_pct": round(
                    100 - zone_metrics["text_area_pct"] - zone_metrics["whitespace_pct"], 2
                ),
                "raw_char_count": tess_metrics["raw_char_count"],
                "processed_char_count": tess_metrics["processed_char_count"],
                "preprocessing_gain_pct": tess_metrics["preprocessing_gain_pct"],
                "accuracy_loss_json": json.dumps(loss["accuracy_loss_breakdown"]),
                "page_metrics_json": "",  # Caller aggregates pages
                "accuracy_tier": self.tier,
            }
        except Exception as exc:
            logger.warning("OCR accuracy analysis failed: %s", exc)
            return _empty_metrics("ocr", self.tier)

    # ------------------------------------------------------------------
    # Layer 1: Pure text (TXT, CSV, MD)
    # ------------------------------------------------------------------
    def _analyze_pure_text(
        self, file_path: str, extracted_text: str, ext: str
    ) -> Dict[str, Any]:
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            m = _empty_metrics("text_extraction", self.tier)
            m["extraction_accuracy"] = 0.0
            m["text_area_pct"] = 0.0
            m["non_text_area_pct"] = 0.0
            return m

        extracted_bytes = len(extracted_text.encode("utf-8"))
        # Text formats: Tika may add/remove whitespace but content is fully extractable
        accuracy = min(extracted_bytes / max(file_size, 1) * 100, 100.0)
        # For CSV, small overhead from delimiters; for TXT/MD essentially 0
        overhead = {".csv": 5.0, ".md": 3.0, ".txt": 0.0}.get(ext, 0.0)

        return {
            "pipeline_type": "text_extraction",
            "extraction_accuracy": round(accuracy, 2),
            "text_area_pct": round(100.0 - overhead, 2),
            "non_text_area_pct": round(overhead, 2),
            "raw_char_count": len(extracted_text),
            "processed_char_count": len(extracted_text),
            "preprocessing_gain_pct": 0.0,
            "accuracy_loss_json": json.dumps(
                {"syntax_overhead_pct": overhead, "format": ext.lstrip(".")}
            ),
            "page_metrics_json": json.dumps(
                [{"page": 1, "text_area_pct": round(100 - overhead, 2), "extraction_accuracy": round(accuracy, 2)}]
            ),
            "accuracy_tier": self.tier,
        }

    # ------------------------------------------------------------------
    # Layer 1b: Structured text (JSON, XML)
    # ------------------------------------------------------------------
    def _analyze_structured_text(
        self, file_path: str, extracted_text: str, ext: str
    ) -> Dict[str, Any]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            raw = f.read()

        if not raw:
            m = _empty_metrics("text_extraction", self.tier)
            m["extraction_accuracy"] = 0.0
            return m

        # Discount structural syntax
        if ext == ".json":
            # Count structural chars: {}[]",:
            struct_chars = sum(1 for c in raw if c in '{}[]",:')
            overhead_pct = struct_chars / max(len(raw), 1) * 100
        else:  # XML
            tags = re.findall(r"<[^>]+>", raw)
            tag_chars = sum(len(t) for t in tags)
            overhead_pct = tag_chars / max(len(raw), 1) * 100

        # Expected text = raw minus structure
        expected_text_chars = len(raw) - int(len(raw) * overhead_pct / 100)
        accuracy = len(extracted_text) / max(expected_text_chars, 1) * 100

        return {
            "pipeline_type": "text_extraction",
            "extraction_accuracy": round(min(accuracy, 100.0), 2),
            "text_area_pct": round(100 - overhead_pct, 2),
            "non_text_area_pct": round(overhead_pct, 2),
            "raw_char_count": expected_text_chars,
            "processed_char_count": len(extracted_text),
            "preprocessing_gain_pct": 0.0,
            "accuracy_loss_json": json.dumps(
                {"syntax_overhead_pct": round(overhead_pct, 2), "format": ext.lstrip(".")}
            ),
            "page_metrics_json": json.dumps(
                [{"page": 1, "text_area_pct": round(100 - overhead_pct, 2), "extraction_accuracy": round(min(accuracy, 100.0), 2)}]
            ),
            "accuracy_tier": self.tier,
        }

    # ------------------------------------------------------------------
    # Layer 2: HTML
    # ------------------------------------------------------------------
    def _analyze_html(self, file_path: str, extracted_text: str) -> Dict[str, Any]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            raw_html = f.read()

        if not raw_html:
            m = _empty_metrics("text_extraction", self.tier)
            m["extraction_accuracy"] = 0.0
            return m

        # Strip all tags to get expected text
        tags = re.findall(r"<[^>]+>", raw_html)
        tag_chars = sum(len(t) for t in tags)
        total_chars = len(raw_html)
        text_area_pct = (total_chars - tag_chars) / max(total_chars, 1) * 100

        # Count embedded images
        img_count = len(re.findall(r"<img[^>]*>", raw_html, re.IGNORECASE))
        image_area_estimate = min(img_count * 5.0, 25.0)

        # Expected text after stripping tags
        expected_text = re.sub(r"<[^>]+>", " ", raw_html)
        expected_text = re.sub(r"\s+", " ", expected_text).strip()
        accuracy = len(extracted_text) / max(len(expected_text), 1) * 100

        return {
            "pipeline_type": "text_extraction",
            "extraction_accuracy": round(min(accuracy, 100.0), 2),
            "text_area_pct": round(text_area_pct, 2),
            "non_text_area_pct": round(100 - text_area_pct, 2),
            "raw_char_count": len(expected_text),
            "processed_char_count": len(extracted_text),
            "preprocessing_gain_pct": 0.0,
            "accuracy_loss_json": json.dumps({
                "html_tags_pct": round(100 - text_area_pct, 2),
                "embedded_images": img_count,
                "image_area_estimate_pct": round(image_area_estimate, 2),
            }),
            "page_metrics_json": json.dumps(
                [{"page": 1, "text_area_pct": round(text_area_pct, 2), "extraction_accuracy": round(min(accuracy, 100.0), 2)}]
            ),
            "accuracy_tier": self.tier,
        }

    # ------------------------------------------------------------------
    # Layer 3a: DOCX
    # ------------------------------------------------------------------
    def _analyze_docx(self, file_path: str, extracted_text: str) -> Dict[str, Any]:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            logger.warning("python-docx not installed — falling back to generic analysis")
            return self._analyze_generic(file_path, extracted_text, ".docx")

        try:
            doc = DocxDocument(file_path)
        except Exception as exc:
            logger.warning("Cannot open DOCX %s: %s", file_path, exc)
            return self._analyze_generic(file_path, extracted_text, ".docx")

        # Count text from paragraphs
        para_chars = sum(len(p.text) for p in doc.paragraphs)

        # Count text from tables
        table_chars = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_chars += len(cell.text)

        total_expected = para_chars + table_chars

        # Count images
        image_count = 0
        try:
            image_count = len(doc.inline_shapes)
        except Exception:
            pass
        # Also count from relationships
        try:
            image_rels = [r for r in doc.part.rels.values()
                          if "image" in str(getattr(r, 'reltype', '')).lower()]
            image_count = max(image_count, len(image_rels))
        except Exception:
            pass

        # Count headers/footers
        hf_chars = 0
        try:
            for section in doc.sections:
                if section.header and section.header.paragraphs:
                    hf_chars += sum(len(p.text) for p in section.header.paragraphs)
                if section.footer and section.footer.paragraphs:
                    hf_chars += sum(len(p.text) for p in section.footer.paragraphs)
        except Exception:
            pass

        total_expected += hf_chars

        # Estimate composition
        if total_expected + image_count > 0:
            text_area_pct = total_expected / (total_expected + image_count * 60) * 100
        else:
            text_area_pct = 100.0

        if total_expected > 0:
            accuracy = min(len(extracted_text) / total_expected * 100, 100.0)
        else:
            accuracy = 100.0 if not extracted_text else 0.0

        return {
            "pipeline_type": "text_extraction",
            "extraction_accuracy": round(accuracy, 2),
            "text_area_pct": round(text_area_pct, 2),
            "non_text_area_pct": round(100 - text_area_pct, 2),
            "raw_char_count": total_expected,
            "processed_char_count": len(extracted_text),
            "preprocessing_gain_pct": 0.0,
            "accuracy_loss_json": json.dumps({
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "table_chars": table_chars,
                "images": image_count,
                "header_footer_chars": hf_chars,
                "image_area_pct": round(100 - text_area_pct, 2),
            }),
            "page_metrics_json": json.dumps(
                [{"page": 1, "text_area_pct": round(text_area_pct, 2),
                  "image_area_pct": round(100 - text_area_pct, 2),
                  "extraction_accuracy": round(accuracy, 2)}]
            ),
            "accuracy_tier": self.tier,
        }

    # ------------------------------------------------------------------
    # Layer 3b: XLSX
    # ------------------------------------------------------------------
    def _analyze_xlsx(self, file_path: str, extracted_text: str) -> Dict[str, Any]:
        try:
            from openpyxl import load_workbook
        except ImportError:
            logger.warning("openpyxl not installed — falling back to generic analysis")
            return self._analyze_generic(file_path, extracted_text, ".xlsx")

        try:
            wb = load_workbook(file_path, data_only=True, read_only=True)
        except Exception as exc:
            logger.warning("Cannot open XLSX %s: %s", file_path, exc)
            return self._analyze_generic(file_path, extracted_text, ".xlsx")

        total_cells = 0
        text_cells = 0
        number_cells = 0
        empty_cells = 0
        text_char_count = 0

        try:
            for ws in wb.worksheets:
                for row in ws.iter_rows():
                    for cell in row:
                        total_cells += 1
                        val = cell.value
                        if val is None or str(val).strip() == "":
                            empty_cells += 1
                        elif isinstance(val, str):
                            text_cells += 1
                            text_char_count += len(val)
                        elif isinstance(val, (int, float)):
                            number_cells += 1
                            text_char_count += len(str(val))
        except Exception as exc:
            logger.warning("Error reading XLSX cells: %s", exc)
        finally:
            try:
                wb.close()
            except Exception:
                pass

        extractable_cells = text_cells + number_cells
        text_area_pct = extractable_cells / max(total_cells, 1) * 100
        accuracy = len(extracted_text) / max(text_char_count, 1) * 100

        return {
            "pipeline_type": "text_extraction",
            "extraction_accuracy": round(min(accuracy, 100.0), 2),
            "text_area_pct": round(text_area_pct, 2),
            "non_text_area_pct": round(100 - text_area_pct, 2),
            "raw_char_count": text_char_count,
            "processed_char_count": len(extracted_text),
            "preprocessing_gain_pct": 0.0,
            "accuracy_loss_json": json.dumps({
                "total_cells": total_cells,
                "text_cells": text_cells,
                "number_cells": number_cells,
                "empty_cells": empty_cells,
                "empty_cell_pct": round(empty_cells / max(total_cells, 1) * 100, 2),
            }),
            "page_metrics_json": json.dumps(
                [{"page": 1, "text_area_pct": round(text_area_pct, 2),
                  "extraction_accuracy": round(min(accuracy, 100.0), 2)}]
            ),
            "accuracy_tier": self.tier,
        }

    # ------------------------------------------------------------------
    # Layer 4a: Text-based PDF
    # ------------------------------------------------------------------
    def _analyze_text_pdf(
        self, file_path: str, extracted_text: str, tika_response: Optional[Dict] = None
    ) -> Dict[str, Any]:
        # Try to get page count from Tika metadata
        page_count = 1
        if tika_response:
            docs = tika_response.get("documents", [])
            if docs:
                meta = docs[0] if isinstance(docs, list) else docs
                page_count = int(
                    meta.get("xmpTPg:NPages", 0) or meta.get("meta:page-count", 0) or 1
                )
        page_count = max(page_count, 1)

        chars_per_page = len(extracted_text) / page_count
        # A typical A4 page has ~2000-3000 chars of text
        EXPECTED_CHARS = 2500
        text_density = min(chars_per_page / EXPECTED_CHARS, 1.0)
        text_area_pct = text_density * 85  # Text max ~85% of page (margins etc)
        non_text_pct = 100 - text_area_pct

        # If very low text density, the PDF might be scanned
        accuracy = min(text_density * 100, 100.0)

        page_metrics = []
        for p in range(1, page_count + 1):
            page_metrics.append({
                "page": p,
                "text_area_pct": round(text_area_pct, 2),
                "non_text_area_pct": round(non_text_pct, 2),
                "extraction_accuracy": round(accuracy, 2),
            })

        return {
            "pipeline_type": "text_extraction",
            "extraction_accuracy": round(accuracy, 2),
            "text_area_pct": round(text_area_pct, 2),
            "non_text_area_pct": round(non_text_pct, 2),
            "raw_char_count": len(extracted_text),
            "processed_char_count": len(extracted_text),
            "preprocessing_gain_pct": 0.0,
            "accuracy_loss_json": json.dumps({
                "page_count": page_count,
                "chars_per_page": round(chars_per_page, 0),
                "text_density": round(text_density, 4),
                "low_text_warning": chars_per_page < 100,
            }),
            "page_metrics_json": json.dumps(page_metrics),
            "accuracy_tier": self.tier,
        }

    # ------------------------------------------------------------------
    # Generic fallback
    # ------------------------------------------------------------------
    def _analyze_generic(
        self, file_path: str, extracted_text: str, ext: str
    ) -> Dict[str, Any]:
        file_size = os.path.getsize(file_path)
        extracted_bytes = len(extracted_text.encode("utf-8"))
        accuracy = min(extracted_bytes / max(file_size, 1) * 100, 100.0) if file_size > 0 else 0.0
        return {
            "pipeline_type": "text_extraction",
            "extraction_accuracy": round(accuracy, 2),
            "text_area_pct": 95.0,
            "non_text_area_pct": 5.0,
            "raw_char_count": len(extracted_text),
            "processed_char_count": len(extracted_text),
            "preprocessing_gain_pct": 0.0,
            "accuracy_loss_json": json.dumps({"format": ext.lstrip("."), "method": "generic_byte_ratio"}),
            "page_metrics_json": "[]",
            "accuracy_tier": self.tier,
        }

    # ==================================================================
    # OCR Internal Methods
    # ==================================================================
    def _segment_page_opencv(self, image_bytes: bytes) -> Dict[str, float]:
        """Segment a page image into text/image/signature/stamp/noise/whitespace zones."""
        nparr = np.frombuffer(image_bytes, np.uint8)
        img = cv2.imdecode(nparr, cv2.IMREAD_GRAYSCALE)
        if img is None:
            return {
                "text_area_pct": 50.0, "image_area_pct": 0.0,
                "signature_area_pct": 0.0, "stamp_area_pct": 0.0,
                "noise_area_pct": 0.0, "whitespace_pct": 50.0,
            }

        h, w = img.shape
        total_pixels = h * w

        # Guard for all-black/very dark or full-bleed pages
        if np.mean(img) < 50:
            return {
                "text_area_pct": 0.0,
                "faded_text_pct": 0.0,
                "image_area_pct": 100.0,
                "signature_area_pct": 0.0,
                "stamp_area_pct": 0.0,
                "noise_area_pct": 0.0,
                "whitespace_pct": 0.0,
                "validated_whitespace_pct": 0.0,
                "bboxes": [],
            }

        # Binarize
        _, binary = cv2.threshold(img, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)

        # ── Pass 1: coarse kernels for text / stamp / logo blocks ─────────────
        # (25,1)+(1,15) merges entire lines together — good for text layout but
        # too aggressive for signatures: the "By:" label + cursive strokes +
        # printed name below all merge into one tall block that fails the
        # height_ratio gate.  Use this pass ONLY for non-signature categories.
        h_kernel_coarse = cv2.getStructuringElement(cv2.MORPH_RECT, (25, 1))
        h_dilated_coarse = cv2.dilate(binary, h_kernel_coarse, iterations=1)
        v_kernel_coarse = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 15))
        text_blocks = cv2.dilate(h_dilated_coarse, v_kernel_coarse, iterations=1)

        coarse_contours, _ = cv2.findContours(
            text_blocks, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE
        )

        # ── Pass 2: tight kernels to isolate signature rows ───────────────────
        # (10,1) bridges the short gaps within a single cursive stroke cluster
        # without pulling in text rows above/below.  (1,4) allows slight vertical
        # variance in ascending/descending loops without merging adjacent lines.
        h_kernel_tight = cv2.getStructuringElement(cv2.MORPH_RECT, (10, 1))
        h_dilated_tight = cv2.dilate(binary, h_kernel_tight, iterations=1)
        v_kernel_tight = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 4))
        sig_blocks = cv2.dilate(h_dilated_tight, v_kernel_tight, iterations=1)

        tight_contours, _ = cv2.findContours(
            sig_blocks, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE
        )

        text_px = 0
        image_px = 0
        sig_px = 0
        stamp_px = 0
        noise_px = 0
        bboxes = []

        # Helper: IoU overlap between two bboxes [x1,y1,x2,y2]
        def _iou(a, b):
            ix1 = max(a[0], b[0]); iy1 = max(a[1], b[1])
            ix2 = min(a[2], b[2]); iy2 = min(a[3], b[3])
            iw = max(0, ix2 - ix1); ih = max(0, iy2 - iy1)
            inter = iw * ih
            if inter == 0:
                return 0.0
            ua = (a[2]-a[0])*(a[3]-a[1]) + (b[2]-b[0])*(b[3]-b[1]) - inter
            return inter / max(ua, 1)

        # ── Process coarse pass — stamp / logo / text (skip signature attempt) ─
        for cnt in coarse_contours:
            x, y, cw, ch = cv2.boundingRect(cnt)
            area = cw * ch
            if area == 0:
                continue
            aspect = cw / max(ch, 1)
            roi = binary[y : y + ch, x : x + cw]
            ink_density = np.count_nonzero(roi) / area

            if area < total_pixels * 0.0005:
                noise_px += area
                continue

            # Stamp/Seal
            if 0.6 < aspect < 1.6 and area > total_pixels * 0.005:
                perimeter = cv2.arcLength(cnt, True)
                if perimeter > 0:
                    circularity = 4 * np.pi * cv2.contourArea(cnt) / (perimeter ** 2)
                    if circularity > 0.4:
                        stamp_px += area
                        bboxes.append({
                            "type": "stamp",
                            "bbox": [int(x), int(y), int(x + cw), int(y + ch)],
                            "impact": round(area / total_pixels * 100, 2)
                        })
                        continue

            # Image/Logo
            if area > total_pixels * 0.02 and area <= total_pixels * 0.30 and ink_density > 0.5:
                image_px += area
                bboxes.append({
                    "type": "logo",
                    "bbox": [int(x), int(y), int(x + cw), int(y + ch)],
                    "impact": round(area / total_pixels * 100, 2)
                })
                continue

            text_px += area

        # ── Process tight pass — signature candidates only ────────────────────
        # Rules are intentionally relaxed compared to the old single-pass:
        #   - position: anywhere in lower 60% (y+ch > h*0.4)
        #   - aspect: 1.5–25  (loopy signatures can be nearly square per row)
        #   - height: 0.002–0.12  (a line-height signature on a short page)
        #   - width: 0.03–0.75  (right-column signatures use less full-page width)
        #   - ink_density: 0.005–0.35  (tight bbox around real strokes is denser)
        for cnt in tight_contours:
            x, y, cw, ch = cv2.boundingRect(cnt)
            area = cw * ch
            if area == 0:
                continue
            if area < total_pixels * 0.0005:
                continue

            aspect = cw / max(ch, 1)
            width_ratio = cw / max(w, 1)
            height_ratio = ch / max(h, 1)
            roi = binary[y : y + ch, x : x + cw]
            ink_density = np.count_nonzero(roi) / area

            if not (
                (y + ch) > h * 0.4
                and 1.5 < aspect < 25.0
                and 0.002 < height_ratio < 0.12
                and 0.03 < width_ratio < 0.75
                and 0.005 < ink_density < 0.35
            ):
                continue

            candidate_bbox = [int(x), int(y), int(x + cw), int(y + ch)]

            # Skip if this box overlaps heavily with an already-found stamp/logo
            overlap = any(
                _iou(candidate_bbox, b["bbox"]) > 0.4
                for b in bboxes
                if b["type"] in ("stamp", "logo")
            )
            if overlap:
                continue

            # Deduplicate against already-accepted signature bboxes
            dup = any(
                _iou(candidate_bbox, b["bbox"]) > 0.5
                for b in bboxes
                if b["type"] == "signature"
            )
            if dup:
                continue

            sig_px += area
            bboxes.append({
                "type": "signature",
                "bbox": candidate_bbox,
                "impact": round(area / total_pixels * 100, 2)
            })

        # Detect faded text regions
        faded_regions = self._detect_faded_text_regions(image_bytes)
        bboxes.extend(faded_regions)
        faded_px = sum((r["bbox"][2] - r["bbox"][0]) * (r["bbox"][3] - r["bbox"][1]) for r in faded_regions)

        covered = text_px + image_px + sig_px + stamp_px + noise_px + faded_px
        ws_px = max(0, total_pixels - covered)

        raw_text_pct     = round(text_px   / total_pixels * 100, 2)
        raw_image_pct    = round(image_px  / total_pixels * 100, 2)
        raw_sig_pct      = round(sig_px    / total_pixels * 100, 2)
        raw_stamp_pct    = round(stamp_px  / total_pixels * 100, 2)
        raw_noise_pct    = round(noise_px  / total_pixels * 100, 2)
        raw_ws_pct       = round(ws_px     / total_pixels * 100, 2)
        raw_faded_pct    = round(faded_px  / total_pixels * 100, 2)

        # Normalize: overlapping contours can cause the sum to exceed 100%.
        # Scale down proportionally so the result always sums to ≤ 100%.
        raw_sum = raw_text_pct + raw_image_pct + raw_sig_pct + raw_stamp_pct + raw_noise_pct + raw_ws_pct + raw_faded_pct
        if raw_sum > 100.0:
            scale = 100.0 / raw_sum
            raw_text_pct  = round(raw_text_pct  * scale, 2)
            raw_image_pct = round(raw_image_pct * scale, 2)
            raw_sig_pct   = round(raw_sig_pct   * scale, 2)
            raw_stamp_pct = round(raw_stamp_pct * scale, 2)
            raw_noise_pct = round(raw_noise_pct * scale, 2)
            raw_ws_pct    = round(raw_ws_pct    * scale, 2)
            raw_faded_pct = round(raw_faded_pct * scale, 2)

        return {
            "text_area_pct":      raw_text_pct,
            "faded_text_pct":     raw_faded_pct,
            "image_area_pct":     raw_image_pct,
            "signature_area_pct": raw_sig_pct,
            "stamp_area_pct":     raw_stamp_pct,
            "noise_area_pct":     raw_noise_pct,
            "whitespace_pct":     raw_ws_pct,
            "validated_whitespace_pct": raw_ws_pct,
            "bboxes": bboxes,
        }


    def _validate_whitespace_region(self, arr: np.ndarray, region: Tuple[int, int, int, int]) -> bool:
        """Validate if a suspect whitespace region is genuinely empty whitespace."""
        x1, y1, x2, y2 = region
        # Clamp coordinates to array bounds
        h_arr, w_arr = arr.shape
        x1 = max(0, min(x1, w_arr - 1))
        y1 = max(0, min(y1, h_arr - 1))
        x2 = max(0, min(x2, w_arr))
        y2 = max(0, min(y2, h_arr))
        
        if x2 <= x1 or y2 <= y1:
            return True
            
        roi = arr[y1:y2, x1:x2]
        # Count ink pixels (pixels darker than 245)
        ink_pixels = np.count_nonzero(roi < 245)
        ink_ratio = ink_pixels / max(1, roi.size)
        return ink_ratio < 0.005

    def _detect_faded_text_regions(self, image_bytes: bytes) -> List[Dict[str, Any]]:
        """Detect faint/faded text regions in the page image."""
        nparr = np.frombuffer(image_bytes, np.uint8)
        img = cv2.imdecode(nparr, cv2.IMREAD_GRAYSCALE)
        if img is None:
            return []
            
        # Guard for all-black or very dark pages (TS-15)
        if np.mean(img) < 50:
            return []
            
        # Binarize with threshold in range [120, 235] to find faint gray text
        mask = cv2.inRange(img, 120, 235)
        
        # Dilate horizontally to merge characters/words, vertically to merge lines close together (TS-13)
        kernel_h = cv2.getStructuringElement(cv2.MORPH_RECT, (50, 1))
        kernel_v = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 25))
        dilated = cv2.dilate(mask, kernel_h, iterations=1)
        dilated = cv2.dilate(dilated, kernel_v, iterations=1)
        
        contours, _ = cv2.findContours(dilated, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        
        faded_regions = []
        h_img, w_img = img.shape
        total_area = h_img * w_img
        
        for cnt in contours:
            x, y, w, h = cv2.boundingRect(cnt)
            area = w * h
            
            # Ignore tiny noise or huge blocks
            if area < total_area * 0.0005 or area > total_area * 0.5:
                continue
                
            roi = img[y:y+h, x:x+w]
            min_val = np.min(roi)
            # Table lines or black text (min_val < 100) are skipped (TS-22, TS-10)
            if min_val < 100:
                continue
                
            # Ink density threshold
            ink_px = np.count_nonzero((roi >= 120) & (roi <= 235))
            if ink_px / max(1, area) < 0.02:
                continue
                
            # Estimate accuracy impact (percent of page content area, default to area percentage)
            impact = round((area / total_area) * 100, 2)
            
            faded_regions.append({
                "type": "faded_text",
                "bbox": [int(x), int(y), int(x + w), int(y + h)],
                "impact": max(0.1, impact)
            })
            
        return faded_regions

    def _refine_zones_yolo(self, image_bytes: bytes, opencv_zones: dict) -> dict:
        """Refine zone detection using YOLO object detection.
        Falls back to opencv_zones if YOLO is not available.
        """
        if not self._yolo:
            return opencv_zones

        try:
            img = Image.open(io.BytesIO(image_bytes))
            w, h = img.size
            total_pixels = w * h

            results = self._yolo(img, conf=0.25, iou=0.45, verbose=False)

            zones = {
                'text_area_pct': 0, 'image_area_pct': 0, 'signature_area_pct': 0,
                'stamp_area_pct': 0, 'table_area_pct': 0, 'noise_area_pct': 0,
            }

            yolo_bboxes = []
            # Generic YOLO class mapping (yolov8n trained on COCO)
            # Map detected objects to document zone types
            for box in results[0].boxes:
                cls_id = int(box.cls)
                x1, y1, x2, y2 = box.xyxy[0].tolist()
                area = (x2 - x1) * (y2 - y1)
                
                # Ignore giant false-positive COCO boxes (e.g. detecting the entire scanned page as a book/laptop/etc.)
                if area > total_pixels * 0.30:
                    continue

                # Most COCO classes are objects/images in document context
                zones['image_area_pct'] += area
                yolo_bboxes.append({
                    "type": "logo",
                    "bbox": [int(x1), int(y1), int(x2), int(y2)],
                    "impact": round(area / total_pixels * 100, 2)
                })

            # Convert to percentages
            for key in zones:
                zones[key] = round(zones[key] / total_pixels * 100, 2)

            # Blend with OpenCV results (YOLO for non-text, OpenCV for text)
            refined = dict(opencv_zones)
            if "bboxes" not in refined:
                refined["bboxes"] = []

            if zones['image_area_pct'] > 0:
                refined['image_area_pct'] = max(opencv_zones.get('image_area_pct', 0), zones['image_area_pct'])
                # Append YOLO bboxes
                for yb in yolo_bboxes:
                    refined["bboxes"].append(yb)
                # Adjust text proportionally
                total_non_text = refined['image_area_pct'] + refined.get('signature_area_pct', 0) + refined.get('stamp_area_pct', 0) + refined.get('noise_area_pct', 0)
                refined['text_area_pct'] = max(0, 100 - total_non_text - refined.get('whitespace_pct', 0))

            refined['detection_method'] = 'opencv+yolo'
            return refined

        except Exception as exc:
            logger.debug("YOLO refinement failed, using OpenCV zones: %s", exc)
            return opencv_zones

    def _measure_tesseract_coverage(
        self, raw_bytes: bytes, processed_bytes: bytes
    ) -> Dict[str, Any]:
        """Run Tesseract TSV on raw vs preprocessed image."""
        if not _HAS_PYTESSERACT:
            return {
                "raw_char_count": 0, "processed_char_count": 0,
                "preprocessing_gain_pct": 0.0,
            }

        def _count_chars(img_bytes: bytes) -> Tuple[int, int]:
            img = Image.open(io.BytesIO(img_bytes))
            try:
                tsv = pytesseract.image_to_data(
                    img, output_type=pytesseract.Output.DICT,
                    config="--oem 1 --psm 3",
                    timeout=30,
                )
            except Exception:
                return 0, 0
            high = 0
            total = 0
            for text, conf in zip(tsv["text"], tsv["conf"]):
                text_str = str(text).strip()
                if not text_str:
                    continue
                c = int(conf)
                total += len(text_str)
                if c > 30:
                    high += len(text_str)
            return high, total

        raw_high, raw_total = _count_chars(raw_bytes)
        proc_high, proc_total = _count_chars(processed_bytes)

        gain = (proc_high - raw_high) / max(raw_high, 1) * 100

        return {
            "raw_char_count": raw_high,
            "raw_total_chars": raw_total,
            "processed_char_count": proc_high,
            "processed_total_chars": proc_total,
            "preprocessing_gain_pct": round(gain, 2),
            "chars_recovered": proc_high - raw_high,
        }

    def _measure_doctr_ground_truth(
        self, raw_bytes: bytes, processed_bytes: bytes
    ) -> Optional[Dict[str, int]]:
        """Use DocTR as ground-truth OCR engine."""
        if not self._doctr:
            return None

        def _count_doctr(img_bytes: bytes) -> int:
            try:
                doc = DocumentFile.from_images([img_bytes])
                result = self._doctr(doc)
                chars = 0
                for page in result.pages:
                    for block in page.blocks:
                        for line in block.lines:
                            for word in line.words:
                                if word.confidence > 0.3:
                                    chars += len(word.value)
                return chars
            except Exception as exc:
                logger.debug("DocTR analysis failed: %s", exc)
                return 0

        raw_chars = _count_doctr(raw_bytes)
        proc_chars = _count_doctr(processed_bytes)

        return {
            "doctr_raw_chars": raw_chars,
            "doctr_processed_chars": proc_chars,
            "doctr_ground_truth": max(raw_chars, proc_chars),
        }

    def _estimate_total_chars(self, image_bytes: bytes, text_area_pct: float) -> int:
        """Estimate total character capacity in the text zone."""
        img = Image.open(io.BytesIO(image_bytes))
        total_px = img.width * img.height
        text_zone_px = total_px * (text_area_pct / 100)
        # At 300 DPI: avg char ~12x18=216px², with spacing ~420px²
        EFFECTIVE_CHAR_AREA = 420
        return max(int(text_zone_px / EFFECTIVE_CHAR_AREA), 1)

    def _build_loss_breakdown(
        self,
        zones: Dict[str, float],
        tess: Dict[str, Any],
        doctr: Optional[Dict[str, int]],
        estimated_total: int,
    ) -> Dict[str, Any]:
        """Build 100%-sum accuracy loss breakdown."""
        text_area = zones["text_area_pct"]

        # Best available ground truth
        if doctr and doctr.get("doctr_ground_truth", 0) > 0:
            ground_truth = doctr["doctr_ground_truth"]
        else:
            ground_truth = estimated_total

        processed_chars = tess.get("processed_char_count", 0)
        read_fraction = min(processed_chars / max(ground_truth, 1), 1.0)

        text_read = text_area * read_fraction
        text_lost = text_area * (1 - read_fraction)

        img_pct = zones.get("image_area_pct", 0)
        sig_pct = zones.get("signature_area_pct", 0)
        stamp_pct = zones.get("stamp_area_pct", 0)
        noise_pct = zones.get("noise_area_pct", 0)
        ws_pct = zones.get("whitespace_pct", 0)

        total = text_read + text_lost + img_pct + sig_pct + stamp_pct + noise_pct + ws_pct
        if abs(total - 100.0) > 0.01:
            ws_pct += 100.0 - total
            ws_pct = max(0, ws_pct)

        return {
            "extraction_accuracy": round(read_fraction * 100, 2),
            "accuracy_loss_breakdown": {
                "text_read_pct": round(text_read, 2),
                "unreadable_text_pct": round(text_lost, 2),
                "logos_images_pct": round(img_pct, 2),
                "signatures_pct": round(sig_pct, 2),
                "stamps_seals_pct": round(stamp_pct, 2),
                "noise_artifacts_pct": round(noise_pct, 2),
                "whitespace_margins_pct": round(ws_pct, 2),
                "snippets": zones.get("bboxes", []),
            },
            "total_accounted": 100.0,
        }
